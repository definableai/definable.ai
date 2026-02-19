"""DOCX parser — extracts text from Word documents using python-docx.

Requires optional dependency: ``python-docx>=1.0.0``
"""

import io
from typing import List, Set

from definable.reader.models import ContentBlock, ReaderConfig
from definable.reader.parsers.base_parser import BaseParser


def _import_docx():
  try:
    import docx

    return docx
  except ImportError as e:
    raise ImportError(
      "DocxParser requires the 'python-docx' package. Install it with: pip install 'python-docx>=1.0.0' or: pip install 'definable[readers]'"
    ) from e


class DocxParser(BaseParser):
  """Parses DOCX files using python-docx (local, no API calls).

  Extracts paragraphs and optionally tables from Word documents.
  """

  def __init__(self, include_tables: bool = True) -> None:
    _import_docx()
    self.include_tables = include_tables

  def supported_mime_types(self) -> List[str]:
    return ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

  def supported_extensions(self) -> Set[str]:
    return {".docx"}

  def parse(self, data: bytes, *, mime_type: str | None = None, config: ReaderConfig | None = None) -> List[ContentBlock]:
    docx = _import_docx()
    doc = docx.Document(io.BytesIO(data))
    blocks: List[ContentBlock] = []

    # Extract paragraphs
    para_parts: List[str] = []
    for para in doc.paragraphs:
      text = para.text.strip()
      if text:
        para_parts.append(text)

    if para_parts:
      blocks.append(
        ContentBlock(
          content_type="text",
          content="\n\n".join(para_parts),
          mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
      )

    # Extract tables
    if self.include_tables:
      for table in doc.tables:
        rows = []
        for row in table.rows:
          cells = [cell.text.strip() for cell in row.cells]
          rows.append("\t".join(cells))
        if rows:
          blocks.append(
            ContentBlock(
              content_type="table",
              content="\n".join(rows),
              mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
          )

    return blocks
